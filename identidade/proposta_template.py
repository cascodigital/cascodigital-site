"""
CASCO DIGITAL — GERADOR DE PROPOSTAS COMERCIAIS
================================================
USO: Copie este arquivo para a pasta do cliente, renomeie para proposta_CLIENTE.py
     e edite apenas a seção DADOS (entre os blocos ===).
     Execute: py proposta_CLIENTE.py
     O PDF será salvo na mesma pasta do script.
"""

import os
import json
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, Image, KeepTogether, PageBreak
)
from reportlab.platypus.flowables import HRFlowable
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# DADOS DO CLIENTE — EDITE AQUI PARA CADA PROPOSTA
# ============================================================

CLIENTE = {
    "nome":     "Nome Da Empresa",
    "cnpj":     "XX.XXX.XXX/0001-XX",
    "contato":  "Nome do Contato",
    "email":    "contato@cliente.com.br",
}

PROPOSTA = {
    "titulo":   "Suporte Técnico Completo — Infraestrutura Híbrida",
    "subtitulo": '"Infraestrutura de ponta sem complexidade desnecessária"',
}
# numero e data são gerados automaticamente (ver _next_numero / _data_hoje)

OBJETO = [
    "Gestão Microsoft 365 (usuários, segurança, compliance)",
    "Manutenção de infraestrutura local (desktops, rede, VPN)",
    "Monitoramento proativo de estabilidade e falhas (24/7)",
    "Proteção de dados (backup em nuvem via Veeam 365)",
    "Assistência técnica ilimitada (helpdesk, troubleshooting, consultoria)",
    "Suporte presencial (2 visitas/mês inclusas)",
]

OBJETO_OBJETIVO = (
    "Operação segura, produtiva e sem surpresas, com um especialista de "
    "referência que conhece cada detalhe do seu ambiente."
)

# Seções do escopo: lista de (titulo_secao, [itens])
ESCOPO = [
    ("Microsoft 365 — Gestão e Segurança", [
        "Administração de usuários, grupos e licenças",
        "Exchange Online: Configuração de caixas, regras de fluxo, DLP básico",
        "Teams: Equipes, canais, políticas de reunião",
        "SharePoint/OneDrive: Organização, permissões, sincronização",
        "Entra ID básico: Sincronização on-premises, MFA, Conditional Access essencial",
        "Monitoramento de compliance e relatórios de segurança",
    ]),
    ("Infraestrutura Local — Desktop e Rede", [
        "Suporte a desktops e notebooks: Troubleshooting, patches, software",
        "Gestão de impressoras e periféricos",
        "Rede local: configuração, monitoramento, regras",
        "Wireless: cobertura, performance, segurança",
        "VPN: Manutenção, acesso remoto seguro, políticas de conectividade",
        "Segurança perimetral: Firewall, detecção de anomalias",
        "Monitoramento proativo 24/7: Estabilidade de serviços, detecção antecipada de falhas, alertas automáticos",
    ]),
    ("Backup e Recuperação de Dados", [
        "Suporte completo ao backup de dados (configuração, monitoramento, restore)",
        "Controle do backup automático de Exchange, SharePoint",
        "Política de retenção (definida conforme necessidade do cliente)",
        "Testes de restore mensais (validação silenciosa)",
        "Consultoria em estratégia de recuperação de desastres",
    ]),
    ("Suporte Técnico Operacional", [
        "Atendimento remoto (chamados ilimitados)",
        "2 visitas presenciais/mês inclusas",
        "Consultoria em hardware (recomendações, orçamentos, fornecedores)",
        "Relatório mensal: Atividades, recomendações, alertas de segurança",
    ]),
    ("Valor Agregado", [
        "Documentação técnica do ambiente (topologia, credenciais seguras, procedimentos)",
        "Consultoria em políticas de segurança para uso de M365",
        "Orientação em adequações de conformidade (LGPD básico)",
    ]),
]

# Tabela de valores: lista de [Item, Detalhes]
VALORES_TABELA = [
    ["Dispositivos gerenciados",    "8"],
    ["Valor Mensal",                "R$ 800,00"],
    ["Atendimento Remoto",          "Chamados ilimitados"],
    ["Atendimento Presencial",      "2 visitas/mês inclusas"],
    ["Visita Adicional",            "R$ 70,00 por visita"],
    ["Horário Padrão",              "Seg-sex, 8h–18h"],
    ["Plantão Remoto",              "R$ 150/h (mín. 1h cheia)"],
    ["Plantão Presencial",          "R$ 150/h (mín. 1h)"],
    ["Pagamento",                   "PIX ou boleto, antecipado até 1º do mês"],
    ["Vigência",                    "12 meses — renovação automática"],
    ["Cancelamento",                "30 dias de aviso prévio"],
]

VALORES_NOTA = (
    "Proposta de Valor: Especialista dedicado em toda sua infraestrutura "
    "(nuvem + local + presencial) com preço justo e sem intermediários."
)

VALORES_PLANTAO = [
    "Plantão é acionamento fora do horário comercial (após 18h, fins de semana, feriados)",
    "Cobrança mínima de 1 hora cheia, depois fracionado em múltiplos de 15 minutos",
    "Plantão presencial: deslocamento conta desde ativação (saída de casa até retorno)",
    "Faturamento de plantão separado, no mês seguinte ao acionamento",
]

# SLA: lista de [Situação, Resposta, Resolução, Prioridade]
SLA_TABELA = [
    ["CRÍTICO — Serviço parado",         "Até 1h",  "Até 4h",  "Máxima"],
    ["ALTO — Funcionalidade impactada",   "Até 2h",  "Até 8h",  "Alta"],
    ["NORMAL — Consultoria, ajustes",     "Até 4h",  "Até 24h", "Padrão"],
]

SLA_NOTA = (
    "SLA em horário comercial (seg-sex, 8h–18h). Atendimentos fora deste "
    "horário configuram plantão (cobrança separada conforme tabela de valores)."
)

EXCLUSOES = [
    "Aquisição ou substituição de hardware (você compra, nós instalamos e configuramos)",
    "Suporte a sistemas legados ou software proprietário específico (caso a caso)",
    "Treinamentos formais estruturados (orçável separadamente)",
    "Alterações de escopo sem comunicação prévia (ex: máquinas adicionais, software novo)",
    "SLA expandido para 24/7 sem plantão (solicitável por adicional)",
    "Visitas presenciais além das 2 mensais (R$ 70 por visita adicional)",
]

CONSIDERACOES = [
    ("Confiável",  "Disponibilidade garantida e SLA respeitado"),
    ("Proativo",   "Prevenção sempre antes da crise"),
    ("Escalável",  "Se crescer para mais máquinas, ajustamos; se reduzir, também"),
    ("Realista",   "Sem promessas impossíveis, apenas o que se entrega"),
]

# ============================================================
# CONFIGURAÇÃO — CASCO DIGITAL (não editar)
# ============================================================

LOGO_PATH = r"C:\Users\kittl\OneDrive - cascodigital\IDENTIDADE\IAs\redondo.png"

CASCO = {
    "responsavel": "André Kittler da Costa",
    "comercial":   "Casco Digital Informática",
    "razao":       "Paloma Arbello Eberhardt",
    "cnpj":        "53.428.468/0001-48",
    "email":       "andre@cascodigital.com.br",
    "celular":     "51 99241-3388",
}

COR_PRIMARIA   = colors.HexColor("#1A5276")   # azul escuro — títulos
COR_SECUNDARIA = colors.HexColor("#2E86C1")   # azul médio — subtítulos
COR_ACENTO     = colors.HexColor("#00D9FF")   # ciano — destaques (identidade site)
COR_TABELA     = colors.HexColor("#D6EAF8")   # azul claro — fundo de tabela
COR_LINHA      = colors.HexColor("#85C1E9")   # azul linha tabela

# ============================================================
# AUTO-NUMERAÇÃO E DATA
# ============================================================

_COUNTER_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".proposta_counter.json")

def _next_numero():
    """Lê e incrementa o contador de propostas. Formato: AAAA-MM-NN (sequência reinicia por mês)."""
    now = datetime.now()
    ano, mes = now.year, now.month
    try:
        with open(_COUNTER_FILE, "r") as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        data = {}
    chave = f"{ano}-{mes:02d}"
    seq = data.get(chave, 0) + 1
    data[chave] = seq
    with open(_COUNTER_FILE, "w") as f:
        json.dump(data, f)
    return f"{ano}-{mes:02d}-{seq:02d}"

def _data_hoje():
    """Retorna data formatada em PT-BR: '5 de março de 2026'."""
    meses = {1:"janeiro",2:"fevereiro",3:"março",4:"abril",5:"maio",6:"junho",
             7:"julho",8:"agosto",9:"setembro",10:"outubro",11:"novembro",12:"dezembro"}
    now = datetime.now()
    return f"{now.day} de {meses[now.month]} de {now.year}"

# ============================================================
# RENDERIZAÇÃO — NÃO EDITAR ABAIXO
# ============================================================

def build_styles():
    base = getSampleStyleSheet()
    styles = {}

    styles["titulo_doc"] = ParagraphStyle("titulo_doc",
        fontSize=20, fontName="Helvetica-Bold",
        textColor=COR_PRIMARIA, alignment=TA_CENTER, spaceAfter=16)

    styles["subtitulo_doc"] = ParagraphStyle("subtitulo_doc",
        fontSize=13, fontName="Helvetica-Bold",
        textColor=COR_SECUNDARIA, alignment=TA_CENTER, spaceAfter=4)

    styles["epigraf"] = ParagraphStyle("epigraf",
        fontSize=10, fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#555555"), alignment=TA_CENTER, spaceAfter=12)

    styles["meta_info"] = ParagraphStyle("meta_info",
        fontSize=9, fontName="Helvetica",
        textColor=colors.HexColor("#666666"), alignment=TA_CENTER, spaceAfter=2)

    styles["cliente_label"] = ParagraphStyle("cliente_label",
        fontSize=9, fontName="Helvetica-Bold",
        textColor=COR_PRIMARIA, spaceAfter=1)

    styles["cliente_valor"] = ParagraphStyle("cliente_valor",
        fontSize=9.5, fontName="Helvetica",
        textColor=colors.HexColor("#222222"), spaceAfter=2, leading=14)

    styles["secao"] = ParagraphStyle("secao",
        fontSize=12, fontName="Helvetica-Bold",
        textColor=COR_PRIMARIA, spaceBefore=14, spaceAfter=6)

    styles["subsecao"] = ParagraphStyle("subsecao",
        fontSize=10, fontName="Helvetica-Bold",
        textColor=COR_SECUNDARIA, spaceBefore=8, spaceAfter=4)

    styles["corpo"] = ParagraphStyle("corpo",
        fontSize=9.5, fontName="Helvetica",
        textColor=colors.HexColor("#222222"), spaceAfter=3, leading=14)

    styles["bullet"] = ParagraphStyle("bullet",
        fontSize=9.5, fontName="Helvetica",
        textColor=colors.HexColor("#222222"),
        leftIndent=18, spaceAfter=7, spaceBefore=2, leading=15,
        bulletIndent=4, bulletFontName="Helvetica", bulletFontSize=9.5)

    styles["nota"] = ParagraphStyle("nota",
        fontSize=9, fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#555555"), spaceAfter=4, leading=13)

    styles["bold_nota"] = ParagraphStyle("bold_nota",
        fontSize=9.5, fontName="Helvetica-Bold",
        textColor=colors.HexColor("#222222"), spaceAfter=6)

    styles["rodape"] = ParagraphStyle("rodape",
        fontSize=8, fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#888888"), alignment=TA_CENTER)

    styles["assinatura_label"] = ParagraphStyle("assinatura_label",
        fontSize=9.5, fontName="Helvetica-Bold",
        textColor=COR_PRIMARIA, spaceAfter=2)

    styles["assinatura_valor"] = ParagraphStyle("assinatura_valor",
        fontSize=9.5, fontName="Helvetica",
        textColor=colors.HexColor("#222222"), spaceAfter=2)

    return styles


def hr(cor=None, espessura=0.5):
    return HRFlowable(width="100%", thickness=espessura,
                      color=cor or COR_LINHA, spaceAfter=14, spaceBefore=4)


def bullet_item(texto, style):
    return Paragraph(f"• {texto}", style)


def secao_com_barra(texto, style):
    """Título de seção com barra ciano à esquerda."""
    t = Table(
        [[Paragraph(texto, style)]],
        colWidths=["100%"],
        spaceBefore=0,
        spaceAfter=8,
    )
    t.setStyle(TableStyle([
        ("LEFTPADDING",   (0, 0), (-1, -1), 10),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LINEBEFORE",    (0, 0), (0, -1), 3, COR_ACENTO),
        ("BACKGROUND",    (0, 0), (-1, -1), colors.HexColor("#F0F8FF")),
    ]))
    return t


def tabela_padrao(dados, col_widths, header=None):
    all_data = ([header] if header else []) + dados
    t = Table(all_data, colWidths=col_widths)
    style = [
        ("FONTNAME",    (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE",    (0, 0), (-1, -1), 9.5),
        ("TOPPADDING",  (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0,0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("GRID",        (0, 0), (-1, -1), 0.4, COR_LINHA),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1),
         [colors.white, COR_TABELA]),
        ("TEXTCOLOR",   (0, 0), (-1, -1), colors.HexColor("#222222")),
        ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
    ]
    if header:
        style += [
            ("BACKGROUND",  (0, 0), (-1, 0), COR_SECUNDARIA),
            ("TEXTCOLOR",   (0, 0), (-1, 0), colors.white),
            ("FONTNAME",    (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",    (0, 0), (-1, 0), 10),
        ]
    t.setStyle(TableStyle(style))
    return t


def gerar_proposta(numero, data_str):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    nome_arquivo = f"Proposta_{CLIENTE['nome'].replace(' ', '_')}.pdf"
    output_path = os.path.join(script_dir, nome_arquivo)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=2.2*cm, rightMargin=2.2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    S = build_styles()
    W = A4[0] - 4.4*cm   # largura útil
    elements = []

    # ── CABEÇALHO — logo + título centralizados ────────────
    logo_w = 2.6*cm
    try:
        logo_img = Image(LOGO_PATH, width=logo_w, height=logo_w)
        logo_img.hAlign = "CENTER"
        elements.append(logo_img)
    except Exception:
        pass
    elements.append(Spacer(1, 6))
    elements.append(Paragraph("PROPOSTA COMERCIAL", S["titulo_doc"]))
    elements.append(Paragraph(PROPOSTA["titulo"], S["subtitulo_doc"]))
    elements.append(Paragraph(PROPOSTA["subtitulo"], S["epigraf"]))
    elements.append(Paragraph(
        f"Proposta nº {numero}  |  {data_str}",
        S["meta_info"]))
    elements.append(hr(COR_PRIMARIA, 1))
    elements.append(Spacer(1, 6))

    # ── CLIENTE (bloco "Para:") ──────────────────────────
    cliente_data = [
        [Paragraph("<b>PARA:</b>", S["cliente_label"]),
         Paragraph(f"<b>{CLIENTE['nome']}</b>", S["cliente_valor"])],
        [Paragraph("CNPJ:", S["cliente_label"]),
         Paragraph(CLIENTE["cnpj"], S["cliente_valor"])],
        [Paragraph("Contato:", S["cliente_label"]),
         Paragraph(f"{CLIENTE['contato']}  —  {CLIENTE['email']}", S["cliente_valor"])],
    ]
    t_cliente = Table(cliente_data, colWidths=[W*0.15, W*0.85])
    t_cliente.setStyle(TableStyle([
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING",   (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 2),
        ("LEFTPADDING",  (0, 0), (-1, -1), 8),
        ("BACKGROUND",   (0, 0), (-1, -1), colors.HexColor("#F5F9FC")),
        ("LINEBEFORE",   (0, 0), (0, -1), 2, COR_ACENTO),
        ("BOX",          (0, 0), (-1, -1), 0.3, COR_LINHA),
    ]))
    elements.append(t_cliente)
    elements.append(Spacer(1, 10))

    # ── OBJETO ─────────────────────────────────────────────
    intro = (f"Prestação de serviços de <b>suporte técnico completo</b> para "
             f"infraestrutura de TI de <b>{CLIENTE['nome']}</b>, cobrindo:")
    elements.append(KeepTogether([
        secao_com_barra("OBJETO", S["secao"]),
        Paragraph(intro, S["corpo"]),
        Spacer(1, 10),
    ]))
    for item in OBJETO:
        elements.append(bullet_item(item, S["bullet"]))
    elements.append(Spacer(1, 6))
    elements.append(Paragraph(f"<b>Objetivo:</b> {OBJETO_OBJETIVO}", S["corpo"]))
    elements.append(hr())

    # ── ESCOPO ─────────────────────────────────────────────
    elements.append(KeepTogether([
        secao_com_barra("ESCOPO DE SERVIÇOS", S["secao"]),
        Paragraph("Atendimento <b>ilimitado</b> para os seguintes componentes:", S["corpo"]),
        Spacer(1, 6),
    ]))
    for titulo_s, itens in ESCOPO:
        bloco = []
        bloco.append(Paragraph(titulo_s, S["subsecao"]))
        bloco.append(Spacer(1, 6))
        for item in itens:
            bloco.append(bullet_item(item, S["bullet"]))
        bloco.append(Spacer(1, 4))
        elements.append(KeepTogether(bloco))
    elements.append(hr())

    # ── VALORES ────────────────────────────────────────────
    valores_data = []
    valor_mensal_row = None
    for idx, r in enumerate(VALORES_TABELA):
        valores_data.append([
            Paragraph(f"<b>{r[0]}</b>", S["corpo"]),
            Paragraph(r[1], S["corpo"]),
        ])
        if "Valor Mensal" in r[0]:
            valor_mensal_row = idx
    t = tabela_padrao(valores_data, col_widths=[W * 0.45, W * 0.55])
    if valor_mensal_row is not None:
        t.setStyle(TableStyle([
            ("BACKGROUND",  (0, valor_mensal_row), (-1, valor_mensal_row), colors.HexColor("#E0F7FA")),
            ("TEXTCOLOR",   (1, valor_mensal_row), (1, valor_mensal_row), colors.HexColor("#0097A7")),
            ("FONTNAME",    (1, valor_mensal_row), (1, valor_mensal_row), "Helvetica-Bold"),
            ("FONTSIZE",    (1, valor_mensal_row), (1, valor_mensal_row), 11),
        ]))
    elements.append(KeepTogether([
        secao_com_barra("VALOR E CONDIÇÕES", S["secao"]),
        t,
    ]))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(f"<b>Proposta de Valor:</b> {VALORES_NOTA.split(':', 1)[-1].strip()}", S["corpo"]))
    elements.append(Spacer(1, 8))
    plantao_bloco = [Paragraph("Observações sobre Plantão", S["subsecao"])]
    for item in VALORES_PLANTAO:
        plantao_bloco.append(bullet_item(item, S["bullet"]))
    elements.append(KeepTogether(plantao_bloco))
    elements.append(hr())

    # ── SLA ────────────────────────────────────────────────
    sla_header = [
        Paragraph("<b>Situação</b>",   S["corpo"]),
        Paragraph("<b>Resposta</b>",   S["corpo"]),
        Paragraph("<b>Resolução</b>",  S["corpo"]),
        Paragraph("<b>Prioridade</b>", S["corpo"]),
    ]
    sla_data = [
        [Paragraph(r[0], S["corpo"]), Paragraph(r[1], S["corpo"]),
         Paragraph(r[2], S["corpo"]), Paragraph(r[3], S["corpo"])]
        for r in SLA_TABELA
    ]
    sla_bloco = [
        secao_com_barra("NÍVEIS DE SERVIÇO (SLA)", S["secao"]),
        Paragraph("Compromisso de resposta conforme criticidade (válido em horário comercial):", S["corpo"]),
        Spacer(1, 6),
        tabela_padrao(sla_data, col_widths=[W*0.45, W*0.18, W*0.18, W*0.19], header=sla_header),
        Spacer(1, 6),
        Paragraph(SLA_NOTA, S["nota"]),
        hr(),
    ]
    elements.append(KeepTogether(sla_bloco))

    # ── EXCLUSÕES ──────────────────────────────────────────
    excl_bloco = [secao_com_barra("ESCOPO DE EXCLUSÃO", S["secao"])]
    excl_bloco.append(Paragraph("Para clareza contratual:", S["corpo"]))
    excl_bloco.append(Spacer(1, 4))
    for item in EXCLUSOES:
        excl_bloco.append(bullet_item(item, S["bullet"]))
    elements.append(KeepTogether(excl_bloco))
    elements.append(hr())

    # ── CONSIDERAÇÕES FINAIS ───────────────────────────────
    consid_bloco = [secao_com_barra("CONSIDERAÇÕES FINAIS", S["secao"])]
    consid_bloco.append(Paragraph(
        "Esta proposta é fruto de análise do seu ambiente atual. Busco ser:", S["corpo"]))
    consid_bloco.append(Spacer(1, 4))
    for chave, valor in CONSIDERACOES:
        consid_bloco.append(Paragraph(f"• <b>{chave}</b> — {valor}", S["bullet"]))
    elements.append(KeepTogether(consid_bloco))
    elements.append(hr())

    # ── ASSINATURA — sempre em página nova ─────────────────
    elements.append(PageBreak())
    elements.append(secao_com_barra("CONTATO E ASSINATURA", S["secao"]))
    elementos_casco = [
        [Paragraph("<b>Responsável técnico:</b>", S["assinatura_label"]),
         Paragraph(CASCO["responsavel"], S["assinatura_valor"])],
        [Paragraph("<b>Nome comercial:</b>", S["assinatura_label"]),
         Paragraph(CASCO["comercial"], S["assinatura_valor"])],
        [Paragraph("<b>Razão social:</b>", S["assinatura_label"]),
         Paragraph(CASCO["razao"], S["assinatura_valor"])],
        [Paragraph("<b>CNPJ:</b>", S["assinatura_label"]),
         Paragraph(CASCO["cnpj"], S["assinatura_valor"])],
        [Paragraph("<b>E-mail:</b>", S["assinatura_label"]),
         Paragraph(CASCO["email"], S["assinatura_valor"])],
        [Paragraph("<b>Celular:</b>", S["assinatura_label"]),
         Paragraph(CASCO["celular"], S["assinatura_valor"])],
    ]
    t_casco = Table(elementos_casco, colWidths=[W*0.35, W*0.65])
    t_casco.setStyle(TableStyle([
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING",  (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING",(0,0), (-1, -1), 2),
    ]))
    elements.append(t_casco)
    elements.append(Spacer(1, 16))

    # bloco de assinaturas
    linha_ass = "_" * 48
    ass_data = [
        [
            Paragraph(f"<b>Pela Contratante ({CLIENTE['nome']}):</b>", S["assinatura_label"]),
            Paragraph("<b>Pela Contratada (Casco Digital Informática):</b>", S["assinatura_label"]),
        ],
        [Spacer(1, 24), Spacer(1, 24)],
        [
            Paragraph(linha_ass, S["nota"]),
            Paragraph(linha_ass, S["nota"]),
        ],
        [
            Paragraph("Nome: ___________________________", S["corpo"]),
            Paragraph(CASCO["responsavel"], S["corpo"]),
        ],
        [
            Paragraph("Cargo: __________________________", S["corpo"]),
            Paragraph(f"Data: {data_str}", S["corpo"]),
        ],
    ]
    t_ass = Table(ass_data, colWidths=[W*0.5, W*0.5])
    t_ass.setStyle(TableStyle([
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING",  (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING",(0,0), (-1, -1), 2),
    ]))
    elements.append(t_ass)
    elements.append(Spacer(1, 16))
    elements.append(Paragraph(
        "Documento confidencial — Uso exclusivo do cliente. Proibida duplicação sem autorização.",
        S["rodape"]))

    doc.build(elements)
    print(f"PDF gerado: {output_path}")


def gerar_docx(numero, data_str):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    nome_arquivo = f"Proposta_{CLIENTE['nome'].replace(' ', '_')}.docx"
    output_path = os.path.join(script_dir, nome_arquivo)

    doc = Document()

    # Margens
    for section in doc.sections:
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    COR_P = RGBColor(0x1A, 0x52, 0x76)
    COR_S = RGBColor(0x2E, 0x86, 0xC1)

    def add_heading(text, level=1, cor=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13 if level == 1 else 10.5)
        run.font.color.rgb = cor or (COR_P if level == 1 else COR_S)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def add_body(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(10)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.space_before = Pt(2)
        return p

    def add_hr():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "85C1E9")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Cabeçalho
    try:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(LOGO_PATH, width=Cm(2.6))
    except Exception:
        pass

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(14)
    r = p_title.add_run("PROPOSTA COMERCIAL")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = COR_P

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run(PROPOSTA["titulo"])
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = COR_S

    p_epi = doc.add_paragraph()
    p_epi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_epi.add_run(PROPOSTA["subtitulo"])
    r3.italic = True
    r3.font.size = Pt(10)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run(f"Proposta nº {numero}  |  {data_str}")
    r_meta.font.size = Pt(9)
    r_meta.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_hr()

    # Cliente
    p_para = doc.add_paragraph()
    r_para = p_para.add_run("PARA: ")
    r_para.bold = True
    r_para.font.size = Pt(9)
    r_para.font.color.rgb = COR_P
    r_nome = p_para.add_run(CLIENTE["nome"])
    r_nome.bold = True
    r_nome.font.size = Pt(10)
    p_para.paragraph_format.space_after = Pt(2)

    p_cnpj = doc.add_paragraph()
    r_cl = p_cnpj.add_run("CNPJ: ")
    r_cl.bold = True
    r_cl.font.size = Pt(9)
    r_cl.font.color.rgb = COR_P
    r_cv = p_cnpj.add_run(CLIENTE["cnpj"])
    r_cv.font.size = Pt(10)
    p_cnpj.paragraph_format.space_after = Pt(2)

    p_cont = doc.add_paragraph()
    r_ctl = p_cont.add_run("Contato: ")
    r_ctl.bold = True
    r_ctl.font.size = Pt(9)
    r_ctl.font.color.rgb = COR_P
    r_ctv = p_cont.add_run(f"{CLIENTE['contato']}  —  {CLIENTE['email']}")
    r_ctv.font.size = Pt(10)
    p_cont.paragraph_format.space_after = Pt(6)

    # Objeto
    add_heading("OBJETO")
    add_body(f"Prestação de serviços de suporte técnico completo para infraestrutura de TI de "
             f"{CLIENTE['nome']}, cobrindo:")
    for item in OBJETO:
        add_bullet(item)
    add_body(f"Objetivo: {OBJETO_OBJETIVO}")
    add_hr()

    # Escopo
    add_heading("ESCOPO DE SERVIÇOS")
    add_body("Atendimento ilimitado para os seguintes componentes:")
    for titulo_s, itens in ESCOPO:
        add_heading(titulo_s, level=2)
        for item in itens:
            add_bullet(item)
    add_hr()

    # Valores
    add_heading("VALOR E CONDIÇÕES")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for row_data in VALORES_TABELA:
        row = table.add_row()
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(9)
    doc.add_paragraph()
    add_body(f"Proposta de Valor: {VALORES_NOTA.split(':', 1)[-1].strip()}", bold=True)
    add_heading("Observações sobre Plantão", level=2)
    for item in VALORES_PLANTAO:
        add_bullet(item)
    add_hr()

    # SLA
    add_heading("NÍVEIS DE SERVIÇO (SLA)")
    add_body("Compromisso de resposta conforme criticidade (válido em horário comercial):")
    sla_table = doc.add_table(rows=1, cols=4)
    sla_table.style = "Table Grid"
    headers = ["Situação", "Resposta", "Resolução", "Prioridade"]
    for i, h in enumerate(headers):
        sla_table.rows[0].cells[i].text = h
        sla_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for row_data in SLA_TABELA:
        row = sla_table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
    doc.add_paragraph()
    p_sla = doc.add_paragraph()
    r_sla = p_sla.add_run(SLA_NOTA)
    r_sla.italic = True
    r_sla.font.size = Pt(9)
    add_hr()

    # Exclusões
    add_heading("ESCOPO DE EXCLUSÃO")
    add_body("Para clareza contratual:")
    for item in EXCLUSOES:
        add_bullet(item)
    add_hr()

    # Considerações
    add_heading("CONSIDERAÇÕES FINAIS")
    add_body("Esta proposta é fruto de análise do seu ambiente atual. Busco ser:")
    for chave, valor in CONSIDERACOES:
        add_bullet(f"{chave} — {valor}")
    add_hr()

    # Assinatura — página nova
    doc.add_page_break()
    add_heading("CONTATO E ASSINATURA")
    for label, valor in [
        ("Responsável técnico", CASCO["responsavel"]),
        ("Nome comercial",      CASCO["comercial"]),
        ("Razão social",        CASCO["razao"]),
        ("CNPJ",                CASCO["cnpj"]),
        ("E-mail",              CASCO["email"]),
        ("Celular",             CASCO["celular"]),
    ]:
        p = doc.add_paragraph()
        r_l = p.add_run(f"{label}: ")
        r_l.bold = True
        r_l.font.size = Pt(10)
        r_v = p.add_run(valor)
        r_v.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    ass_table = doc.add_table(rows=4, cols=2)
    ass_table.style = "Table Grid"
    ass_table.rows[0].cells[0].text = f"Pela Contratante ({CLIENTE['nome']}):"
    ass_table.rows[0].cells[1].text = "Pela Contratada (Casco Digital Informática):"
    for c in [0, 1]:
        ass_table.rows[0].cells[c].paragraphs[0].runs[0].bold = True
    ass_table.rows[1].cells[0].text = " " * 60
    ass_table.rows[1].cells[1].text = " " * 60
    ass_table.rows[2].cells[0].text = "Nome: "
    ass_table.rows[2].cells[1].text = CASCO["responsavel"]
    ass_table.rows[3].cells[0].text = "Cargo: "
    ass_table.rows[3].cells[1].text = f"Data: {data_str}"
    for row in ass_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    p_conf = doc.add_paragraph(
        "Documento confidencial — Uso exclusivo do cliente. Proibida duplicação sem autorização.")
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_conf.runs[0].font.size = Pt(8)
    p_conf.runs[0].italic = True

    doc.save(output_path)
    print(f"DOCX gerado: {output_path}")


if __name__ == "__main__":
    numero = _next_numero()
    data_str = _data_hoje()
    print(f"Proposta nº {numero} — {data_str}")
    gerar_proposta(numero, data_str)
    gerar_docx(numero, data_str)
